"""CV text extraction from PDF and DOCX bytes."""

from __future__ import annotations

import io
import re

from backend_shared.exceptions import ValidationError
from pypdf import PdfReader


def clean_extracted_text(text: str) -> str:
    """Normalize whitespace while preserving paragraph breaks."""
    # Normalize newlines
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Collapse runs of spaces/tabs
    text = re.sub(r"[ \t]+", " ", text)
    # Collapse 3+ blank lines to 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text_from_pdf(data: bytes) -> str:
    """Extract text from a PDF document."""
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ValidationError(
            "Could not read PDF file",
            details={"reason": str(exc)},
        ) from exc

    parts: list[str] = []
    for page in reader.pages:
        try:
            page_text = page.extract_text() or ""
        except Exception:  # noqa: BLE001
            page_text = ""
        if page_text:
            parts.append(page_text)
    return clean_extracted_text("\n\n".join(parts))


def extract_text_from_docx(data: bytes) -> str:
    """Extract text from a DOCX document."""
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise ValidationError("DOCX support is not installed") from exc

    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ValidationError(
            "Could not read DOCX file",
            details={"reason": str(exc)},
        ) from exc

    parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    # Tables
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return clean_extracted_text("\n".join(parts))


def extract_text(data: bytes, *, filename: str = "", content_type: str = "") -> str:
    """Dispatch extraction based on filename extension or content type."""
    name = filename.lower()
    ct = (content_type or "").lower()

    if name.endswith(".pdf") or "pdf" in ct:
        text = extract_text_from_pdf(data)
    elif name.endswith(".docx") or "wordprocessingml" in ct or name.endswith(".doc"):
        if name.endswith(".doc") and not name.endswith(".docx"):
            raise ValidationError(
                "Legacy .doc format is not supported; please upload PDF or DOCX",
                details={"field": "file"},
            )
        text = extract_text_from_docx(data)
    else:
        # Sniff PDF magic
        if data[:4] == b"%PDF":
            text = extract_text_from_pdf(data)
        else:
            raise ValidationError(
                "Unsupported CV format for extraction",
                details={"filename": filename, "content_type": content_type},
            )

    if not text:
        raise ValidationError(
            "No extractable text found in CV",
            details={"filename": filename},
        )
    return text
