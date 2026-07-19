"""CV extraction + processing job tests (Phase 1 PR-009)."""

from __future__ import annotations

import io

import boto3
import pytest
from docx import Document
from fastapi.testclient import TestClient
from moto import mock_aws

from src.profiles.extraction import clean_extracted_text, extract_text, extract_text_from_docx
from src.profiles.storage import get_s3_client
from src.settings import get_settings

STRONG_PASSWORD = "SecurePass1!"


def _docx_bytes(text: str = "Jane Doe\nSenior Engineer\nSkills: Python, FastAPI") -> bytes:
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _register_login(client: TestClient, email: str) -> str:
    client.post(
        "/auth/register",
        json={
            "email": email,
            "password": STRONG_PASSWORD,
            "first_name": "J",
            "last_name": "D",
        },
    )
    login = client.post(
        "/auth/login",
        json={"email": email, "password": STRONG_PASSWORD},
    )
    assert login.status_code == 200, login.text
    return login.json()["data"]["access_token"]


@pytest.fixture()
def s3_moto(monkeypatch: pytest.MonkeyPatch):
    monkeypatch.setenv("AWS_ENDPOINT_URL", "")
    get_settings.cache_clear()
    get_s3_client.cache_clear()
    with mock_aws():
        settings = get_settings()
        client = boto3.client(
            "s3",
            region_name=settings.aws_default_region,
            aws_access_key_id=settings.aws_access_key_id,
            aws_secret_access_key=settings.aws_secret_access_key,
        )
        client.create_bucket(Bucket=settings.s3_bucket)
        get_s3_client.cache_clear()
        yield client
        get_s3_client.cache_clear()
        get_settings.cache_clear()


def test_clean_extracted_text() -> None:
    cleaned = clean_extracted_text("  a \t b  \n\n\n c ")
    assert "a b" in cleaned
    assert "c" in cleaned
    assert "\n\n\n" not in cleaned


def test_extract_docx() -> None:
    data = _docx_bytes("Hello Twin\nWorld")
    text = extract_text_from_docx(data)
    assert "Hello Twin" in text
    assert "World" in text


def test_extract_text_dispatch_docx() -> None:
    data = _docx_bytes("Dispatch OK")
    assert "Dispatch OK" in extract_text(data, filename="cv.docx")


def test_extract_empty_raises() -> None:
    data = _docx_bytes("")
    # Empty paragraphs only → no text
    doc = Document()
    buf = io.BytesIO()
    doc.save(buf)
    with pytest.raises(Exception):
        extract_text(buf.getvalue(), filename="empty.docx")


def test_process_cv_happy_path(client: TestClient, s3_moto) -> None:
    token = _register_login(client, "cv-job@example.com")
    headers = {"Authorization": f"Bearer {token}"}
    data = _docx_bytes("Alex Example\nStaff Engineer\nPython FastAPI PostgreSQL")

    up = client.post(
        "/profiles/me/cv",
        headers=headers,
        files={
            "file": (
                "resume.docx",
                io.BytesIO(data),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert up.status_code == 201, up.text

    proc = client.post("/profiles/me/process-cv", headers=headers)
    assert proc.status_code == 202, proc.text
    job = proc.json()["data"]
    assert job["status"] in {"pending", "processing", "completed"}
    job_id = job["id"]

    status = client.get(f"/profiles/me/cv/jobs/{job_id}", headers=headers)
    assert status.status_code == 200, status.text
    # Eager Celery should have completed
    assert status.json()["data"]["status"] == "completed"
    assert status.json()["data"]["error_message"] is None

    me = client.get("/profiles/me", headers=headers)
    assert me.json()["data"]["has_extracted_text"] is True


def test_process_cv_without_upload(client: TestClient) -> None:
    token = _register_login(client, "cv-no-file@example.com")
    response = client.post(
        "/profiles/me/process-cv",
        headers={"Authorization": f"Bearer {token}"},
    )
    assert response.status_code == 422


def test_process_cv_failed_extraction(client: TestClient, s3_moto) -> None:
    token = _register_login(client, "cv-fail@example.com")
    headers = {"Authorization": f"Bearer {token}"}
    # Valid PDF header but no extractable text → extraction ValidationError
    bad_pdf = b"%PDF-1.4\n1 0 obj<<>>endobj\ntrailer<<>>\n%%EOF\n"
    up = client.post(
        "/profiles/me/cv",
        headers=headers,
        files={"file": ("blank.pdf", io.BytesIO(bad_pdf), "application/pdf")},
    )
    assert up.status_code == 201, up.text

    proc = client.post("/profiles/me/process-cv", headers=headers)
    assert proc.status_code == 202, proc.text
    job_id = proc.json()["data"]["id"]
    status = client.get(f"/profiles/me/cv/jobs/{job_id}", headers=headers)
    assert status.json()["data"]["status"] == "failed"
    assert status.json()["data"]["error_message"]


def test_job_not_found(client: TestClient) -> None:
    token = _register_login(client, "cv-missing-job@example.com")
    import uuid

    response = client.get(
        f"/profiles/me/cv/jobs/{uuid.uuid4()}",
        headers={"Authorization": f"Bearer {token}"},
    )
    assert response.status_code == 404
